#!/usr/bin/env python3
"""将 markdown 书稿转换为带页码目录的 Word 文档"""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, name='SimSun', size=Pt(11), bold=False):
    run.font.size = size
    run.bold = bold
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.insert(0, rFonts)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, 'SimHei', Pt(22), bold=True)
        p.space_after = Pt(6)
    elif level == 1:
        run = p.add_run(text)
        set_run_font(run, 'SimHei', Pt(16), bold=True)
        p.space_before = Pt(18)
        p.space_after = Pt(8)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, 'SimHei', Pt(13), bold=True)
        p.space_before = Pt(12)
        p.space_after = Pt(6)
    elif level == 3:
        run = p.add_run(text)
        set_run_font(run, 'SimHei', Pt(11), bold=True)
        p.space_before = Pt(8)
        p.space_after = Pt(4)
    else:
        run = p.add_run(text)
        set_run_font(run, 'SimSun', Pt(10.5))
    return p

def add_body(doc, text):
    """Add body text, handling bold markers and inline code"""
    p = doc.add_paragraph()
    p.space_after = Pt(4)
    p.space_before = Pt(2)
    # Split by bold ** markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, 'SimHei', Pt(10.5), bold=True)
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            set_run_font(run, 'Consolas', Pt(10))
        else:
            run = p.add_run(part)
            set_run_font(run, 'SimSun', Pt(10.5))
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.space_after = Pt(2)
    p.space_before = Pt(2)
    run = p.add_run(code_text)
    set_run_font(run, 'Consolas', Pt(9))
    # Light gray background via shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    p.paragraph_format.element.get_or_add_pPr().append(shading)

def add_table_from_md(doc, lines):
    """Add a simple table from markdown table lines"""
    if len(lines) < 2:
        return
    # Parse headers
    headers = [h.strip() for h in lines[0].strip('|').split('|')]
    ncols = len(headers)
    # Skip separator line
    data_lines = [l for l in lines[2:] if l.strip().startswith('|')]
    
    table = doc.add_table(rows=1 + len(data_lines), cols=ncols)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, 'SimHei', Pt(9), bold=True)
    
    # Data rows
    for r, line in enumerate(data_lines):
        cells = [c.strip() for c in line.strip('|').split('|')]
        for c, val in enumerate(cells):
            if c < ncols:
                cell = table.rows[r + 1].cells[c]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                set_run_font(run, 'SimSun', Pt(9))

def add_checklist(doc, items):
    for item in items:
        if item.strip():
            p = doc.add_paragraph()
            p.space_after = Pt(2)
            run = p.add_run('☐ ' + item.lstrip('- [ ] ').lstrip('- '))
            set_run_font(run, 'SimSun', Pt(10.5))

def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.space_after = Pt(6)
    run = p.add_run(text.lstrip('> '))
    set_run_font(run, 'KaiTi', Pt(10.5))
    run.italic = True

def add_hr(doc):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Main ──

md_path = 'docs/the-philosophy-of-subtraction.md'
doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(14.8)
    section.page_height = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# We'll track TOC entries with placeholder page numbers
toc_entries = []  # (level, title, paragraph_obj)

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_code_block = False
code_buffer = []
in_table = False
table_buffer = []
in_checklist = False
checklist_buffer = []
skip_until_after_hr = False  # skip TOC section in body

while i < len(lines):
    line = lines[i].rstrip()
    
    # Skip the markdown TOC section (between first --- and next ---)
    if line == '---' and i < 30:
        skip_until_after_hr = not skip_until_after_hr
        if skip_until_after_hr:
            add_hr(doc)
        i += 1
        continue
    
    # Skip TOC lines
    if skip_until_after_hr:
        if line.startswith('- [') or line.startswith('**') or line == '' or line.startswith('>'):
            i += 1
            continue
    
    # Code blocks
    if line.startswith('```'):
        if in_code_block:
            for cl in code_buffer:
                add_code_block(doc, cl)
            code_buffer = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue
    
    if in_code_block:
        code_buffer.append(line)
        i += 1
        continue
    
    # Tables
    if line.startswith('|') and not in_table:
        in_table = True
        table_buffer = [line]
        i += 1
        continue
    if in_table:
        if line.startswith('|'):
            table_buffer.append(line)
            i += 1
            continue
        else:
            add_table_from_md(doc, table_buffer)
            table_buffer = []
            in_table = False
            # Don't increment i - process this line
    
    # Checklists
    if line.startswith('- [ ]'):
        if not in_checklist:
            in_checklist = True
            checklist_buffer = [line]
        else:
            checklist_buffer.append(line)
        i += 1
        continue
    if in_checklist and not line.startswith('- [ ]'):
        add_checklist(doc, checklist_buffer)
        checklist_buffer = []
        in_checklist = False
        # Don't increment i
    
    # ASCII art blocks (the "一条线" diagram)
    if line.strip().startswith('┌') or line.strip().startswith('│') or line.strip().startswith('└'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, 'Consolas', Pt(8))
        p.space_after = Pt(0)
        p.space_before = Pt(0)
        i += 1
        continue
    
    # Headings
    if line.startswith('# '):
        add_heading_styled(doc, line[2:], 0)
        toc_entries.append((0, line[2:], None))
    elif line.startswith('## '):
        p = add_heading_styled(doc, line[3:], 1)
        toc_entries.append((1, line[3:], p))
    elif line.startswith('### '):
        add_heading_styled(doc, line[4:], 2)
    elif line.startswith('#### '):
        add_heading_styled(doc, line[5:], 3)
    
    # Blockquotes
    elif line.startswith('> '):
        add_blockquote(doc, line)
    
    # Horizontal rules
    elif line.strip() == '---':
        add_hr(doc)
    
    # Empty lines
    elif line.strip() == '':
        pass  # skip
    
    # Body text
    elif line.strip():
        add_body(doc, line)
    
    i += 1

# Build TOC page (insert at beginning)
# We need to save first, then insert TOC
output_path = 'docs/the-philosophy-of-subtraction.docx'
doc.save(output_path)

# Now rebuild with TOC at the front
doc2 = Document()
for section in doc2.sections:
    section.page_width = Cm(14.8)
    section.page_height = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Title page
p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run('减法的哲学')
set_run_font(run, 'SimHei', Pt(26), bold=True)

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('The Philosophy of Subtraction')
set_run_font(run, 'SimSun', Pt(14))

p = doc2.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(24)
run = p.add_run('一个不会写代码的人，用 AI 开发 AI 的 69 天实验')
set_run_font(run, 'KaiTi', Pt(12))

doc2.add_page_break()

# TOC page
p = doc2.add_paragraph()
run = p.add_run('目  录')
set_run_font(run, 'SimHei', Pt(18), bold=True)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_after = Pt(18)

# Group TOC by section
sections_toc = {
    '引子': [],
    '间奏：心路历程': [],
    '第一章': [], '第二章': [], '第三章': [], '第四章': [], '第五章': [],
    '第六章': [], '第七章': [], '第八章': [], '第九章': [], '第十章': [],
    '第十一章：数据的证词': [],
    '第十二章：一条线': [],
    '后记': [],
    '附录': [],
}

for level, title, para in toc_entries:
    # Assign page numbers
    page_num = "—"
    if title == '引子':
        page_num = "1"
        sections_toc['引子'].append((level, title, page_num))
    elif title == '间奏：心路历程':
        page_num = "3"
        sections_toc['间奏：心路历程'].append((level, title, page_num))
    elif title.startswith('第一章'):
        page_num = "8"
        sections_toc['第一章'].append((level, title, page_num))
    elif title.startswith('第二章'):
        page_num = "14"
        sections_toc['第二章'].append((level, title, page_num))
    elif title.startswith('第三章'):
        page_num = "18"
        sections_toc['第三章'].append((level, title, page_num))
    elif title.startswith('第四章'):
        page_num = "22"
        sections_toc['第四章'].append((level, title, page_num))
    elif title.startswith('第五章'):
        page_num = "26"
        sections_toc['第五章'].append((level, title, page_num))
    elif title.startswith('第六章'):
        page_num = "29"
        sections_toc['第六章'].append((level, title, page_num))
    elif title.startswith('第七章'):
        page_num = "32"
        sections_toc['第七章'].append((level, title, page_num))
    elif title.startswith('第八章'):
        page_num = "35"
        sections_toc['第八章'].append((level, title, page_num))
    elif title.startswith('第九章'):
        page_num = "38"
        sections_toc['第九章'].append((level, title, page_num))
    elif title.startswith('第十章'):
        page_num = "40"
        sections_toc['第十章'].append((level, title, page_num))
    elif title.startswith('第十一章'):
        page_num = "43"
        sections_toc['第十一章：数据的证词'].append((level, title, page_num))
    elif title.startswith('第十二章'):
        page_num = "47"
        sections_toc['第十二章：一条线'].append((level, title, page_num))
    elif title.startswith('后记'):
        page_num = "50"
        sections_toc['后记'].append((level, title, page_num))
    elif title.startswith('附录'):
        page_num = "52"
        sections_toc['附录'].append((level, title, page_num))

def add_toc_line(doc2, title, page, indent=0):
    p = doc2.add_paragraph()
    p.space_after = Pt(4)
    prefix = '    ' * indent
    run = p.add_run(f'{prefix}{title}')
    set_run_font(run, 'SimSun', Pt(10.5))
    # Add dots and page number
    run2 = p.add_run(f'  {"·" * (40 - len(title) - indent * 4)} {page}')
    set_run_font(run2, 'SimSun', Pt(9))

# Render TOC
order = ['引子', '间奏：心路历程',
         '第一章', '第二章', '第三章', '第四章', '第五章',
         '第六章', '第七章', '第八章', '第九章', '第十章',
         '第十一章：数据的证词', '第十二章：一条线',
         '后记', '附录']

section_labels = {
    '引子': ('', ''),
    '间奏：心路历程': ('第一部分  这条路怎么走过来的', ''),
    '第一章': ('第二部分  十条工程学原则', ''),
    '第十一章：数据的证词': ('第三部分  证据与全貌', ''),
    '后记': ('', ''),
    '附录': ('', ''),
}

for key in order:
    if key in section_labels and section_labels[key][0]:
        p = doc2.add_paragraph()
        p.space_before = Pt(10)
        run = p.add_run(section_labels[key][0])
        set_run_font(run, 'SimHei', Pt(11), bold=True)
    if key in sections_toc:
        for _, title, page in sections_toc[key]:
            indent = 1 if key in ['第一章','第二章','第三章','第四章','第五章','第六章','第七章','第八章','第九章','第十章'] else 0
            add_toc_line(doc2, title, page, indent)

doc2.add_page_break()

# Copy all body content from first doc
for element in doc.element.body:
    # Skip the first heading (it's the title, we already have title page)
    doc2.element.body.append(element)

doc2.save(output_path)
print(f'Word document saved to {output_path}')
